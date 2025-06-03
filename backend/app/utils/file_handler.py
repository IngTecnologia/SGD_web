"""
Utilidades para manejo de archivos
Upload, validación, procesamiento y gestión de archivos
"""
import logging
import os
import shutil
import hashlib
import mimetypes
from typing import List, Dict, Any, Optional, Tuple, BinaryIO
from datetime import datetime
from pathlib import Path
import tempfile

from fastapi import UploadFile, HTTPException, status
from PIL import Image, ExifTags
import fitz  # PyMuPDF
import magic

from ..config import get_settings
from ..schemas.document import DocumentFileInfo

# Configuración
settings = get_settings()
logger = logging.getLogger(__name__)


class FileHandlerError(Exception):
    """Excepción personalizada para errores de manejo de archivos"""
    pass


class FileHandler:
    """
    Clase principal para manejo de archivos
    Validación, procesamiento y almacenamiento seguro
    """
    
    def __init__(self):
        """Inicializar el manejador de archivos"""
        self.documents_path = settings.DOCUMENTS_PATH
        self.temp_path = settings.TEMP_PATH
        self.max_file_size = settings.MAX_FILE_SIZE
        self.allowed_types = settings.ALLOWED_FILE_TYPES
        self.allowed_extensions = settings.ALLOWED_FILE_EXTENSIONS
        
        # Crear directorios si no existen
        os.makedirs(self.documents_path, exist_ok=True)
        os.makedirs(self.temp_path, exist_ok=True)
        
        # Configurar magic para detección de tipos MIME
        try:
            self.magic_mime = magic.Magic(mime=True)
        except:
            logger.warning("python-magic no disponible, usando fallback")
            self.magic_mime = None
        
        logger.info("FileHandler inicializado correctamente")
    
    # === VALIDACIÓN DE ARCHIVOS ===
    
    def validate_file(self, file: UploadFile, allowed_types: List[str] = None) -> Tuple[bool, str]:
        """
        Validar archivo subido
        
        Args:
            file: Archivo subido
            allowed_types: Tipos MIME permitidos (usa configuración por defecto si no se especifica)
            
        Returns:
            Tuple[bool, str]: (es_válido, mensaje_error)
        """
        try:
            if not file:
                return False, "No se proporcionó archivo"
            
            if not file.filename:
                return False, "Nombre de archivo vacío"
            
            # Validar tamaño
            if hasattr(file, 'size') and file.size:
                if file.size > self.max_file_size:
                    max_mb = self.max_file_size // (1024 * 1024)
                    return False, f"Archivo demasiado grande. Máximo: {max_mb}MB"
            
            # Validar extensión
            file_ext = os.path.splitext(file.filename)[1].lower()
            if file_ext not in self.allowed_extensions:
                return False, f"Extensión no permitida: {file_ext}"
            
            # Validar tipo MIME si se especifica
            if allowed_types is None:
                allowed_types = self.allowed_types
            
            # Leer una pequeña porción para detectar tipo MIME
            file_header = file.file.read(8192)
            file.file.seek(0)  # Resetear posición
            
            detected_mime = self._detect_mime_type(file_header, file.filename)
            
            if detected_mime not in allowed_types:
                return False, f"Tipo de archivo no permitido: {detected_mime}"
            
            # Validaciones específicas por tipo
            validation_result = self._validate_file_content(file_header, detected_mime, file.filename)
            if not validation_result[0]:
                return validation_result
            
            return True, "Archivo válido"
            
        except Exception as e:
            logger.error(f"Error validando archivo: {str(e)}")
            return False, f"Error validando archivo: {str(e)}"
    
    def _detect_mime_type(self, file_header: bytes, filename: str) -> str:
        """
        Detectar tipo MIME del archivo
        
        Args:
            file_header: Primeros bytes del archivo
            filename: Nombre del archivo
            
        Returns:
            str: Tipo MIME detectado
        """
        try:
            # Usar python-magic si está disponible
            if self.magic_mime:
                try:
                    return self.magic_mime.from_buffer(file_header)
                except:
                    pass
            
            # Fallback usando mimetypes
            mime_type, _ = mimetypes.guess_type(filename)
            if mime_type:
                return mime_type
            
            # Detección manual basada en headers
            if file_header.startswith(b'%PDF'):
                return 'application/pdf'
            elif file_header.startswith(b'\xff\xd8\xff'):
                return 'image/jpeg'
            elif file_header.startswith(b'\x89PNG\r\n\x1a\n'):
                return 'image/png'
            elif file_header.startswith(b'PK\x03\x04') and filename.endswith('.docx'):
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif file_header.startswith(b'\xd0\xcf\x11\xe0') and filename.endswith('.doc'):
                return 'application/msword'
            else:
                return 'application/octet-stream'
                
        except Exception as e:
            logger.warning(f"Error detectando tipo MIME: {str(e)}")
            return 'application/octet-stream'
    
    def _validate_file_content(self, file_header: bytes, mime_type: str, filename: str) -> Tuple[bool, str]:
        """
        Validar contenido específico del archivo
        
        Args:
            file_header: Primeros bytes del archivo
            mime_type: Tipo MIME detectado
            filename: Nombre del archivo
            
        Returns:
            Tuple[bool, str]: (es_válido, mensaje_error)
        """
        try:
            # Validar PDFs
            if mime_type == 'application/pdf':
                if not file_header.startswith(b'%PDF'):
                    return False, "Archivo PDF corrupto o inválido"
            
            # Validar imágenes JPEG
            elif mime_type == 'image/jpeg':
                if not file_header.startswith(b'\xff\xd8\xff'):
                    return False, "Archivo JPEG corrupto o inválido"
            
            # Validar imágenes PNG
            elif mime_type == 'image/png':
                if not file_header.startswith(b'\x89PNG\r\n\x1a\n'):
                    return False, "Archivo PNG corrupto o inválido"
            
            # Validar documentos Word
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                # Verificar que tenga estructura válida
                if not (file_header.startswith(b'PK\x03\x04') or file_header.startswith(b'\xd0\xcf\x11\xe0')):
                    return False, "Documento Word corrupto o inválido"
            
            # Verificar que no sea un archivo ejecutable disfrazado
            dangerous_headers = [
                b'MZ',  # Ejecutables Windows
                b'\x7fELF',  # Ejecutables Linux
                b'\xfe\xed\xfa',  # Ejecutables macOS
                b'#!/bin/',  # Scripts shell
                b'<?php',  # Scripts PHP
            ]
            
            for dangerous_header in dangerous_headers:
                if file_header.startswith(dangerous_header):
                    return False, "Archivo potencialmente peligroso detectado"
            
            return True, "Contenido válido"
            
        except Exception as e:
            logger.warning(f"Error validando contenido: {str(e)}")
            return True, "Validación de contenido omitida"
    
    # === PROCESAMIENTO DE ARCHIVOS ===
    
    async def process_upload(
        self, 
        file: UploadFile, 
        document_type_id: int,
        user_id: int,
        custom_filename: str = None
    ) -> DocumentFileInfo:
        """
        Procesar archivo subido
        
        Args:
            file: Archivo subido
            document_type_id: ID del tipo de documento
            user_id: ID del usuario que sube
            custom_filename: Nombre personalizado (opcional)
            
        Returns:
            DocumentFileInfo: Información del archivo procesado
            
        Raises:
            FileHandlerError: Si hay errores en el procesamiento
        """
        try:
            logger.info(f"Procesando archivo: {file.filename}")
            
            # Validar archivo
            is_valid, error_msg = self.validate_file(file)
            if not is_valid:
                raise FileHandlerError(f"Archivo inválido: {error_msg}")
            
            # Leer contenido del archivo
            file_content = await file.read()
            file_size = len(file_content)
            
            # Generar hash del archivo
            file_hash = self._calculate_file_hash(file_content)
            
            # Detectar tipo MIME
            mime_type = self._detect_mime_type(file_content[:8192], file.filename)
            
            # Generar nombre único para almacenamiento
            storage_filename = self._generate_storage_filename(
                original_filename=custom_filename or file.filename,
                document_type_id=document_type_id,
                user_id=user_id,
                file_hash=file_hash[:8]
            )
            
            # Crear estructura de carpetas por fecha
            date_path = datetime.now().strftime("%Y/%m/%d")
            storage_dir = os.path.join(self.documents_path, date_path)
            os.makedirs(storage_dir, exist_ok=True)
            
            # Ruta completa del archivo
            storage_path = os.path.join(storage_dir, storage_filename)
            
            # Guardar archivo
            with open(storage_path, 'wb') as f:
                f.write(file_content)
            
            # Crear información del archivo
            file_info = DocumentFileInfo(
                file_name=file.filename,
                file_size=file_size,
                mime_type=mime_type,
                file_hash=file_hash
            )
            
            # Procesar metadatos específicos del tipo de archivo
            metadata = await self._extract_file_metadata(storage_path, mime_type)
            
            logger.info(f"Archivo procesado exitosamente: {storage_filename}")
            
            return {
                "file_info": file_info,
                "storage_path": os.path.relpath(storage_path, self.documents_path),
                "full_path": storage_path,
                "metadata": metadata
            }
            
        except FileHandlerError:
            raise
        except Exception as e:
            logger.error(f"Error procesando archivo: {str(e)}")
            raise FileHandlerError(f"Error procesando archivo: {str(e)}")
    
    def _generate_storage_filename(
        self, 
        original_filename: str, 
        document_type_id: int,
        user_id: int,
        file_hash: str
    ) -> str:
        """
        Generar nombre único para almacenamiento
        
        Args:
            original_filename: Nombre original del archivo
            document_type_id: ID del tipo de documento
            user_id: ID del usuario
            file_hash: Hash parcial del archivo
            
        Returns:
            str: Nombre único para almacenamiento
        """
        try:
            # Extraer extensión
            _, ext = os.path.splitext(original_filename)
            
            # Limpiar nombre original
            clean_name = self._sanitize_filename(
                os.path.splitext(original_filename)[0]
            )
            
            # Generar timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            
            # Construir nombre único
            storage_name = f"doc_{document_type_id}_{user_id}_{timestamp}_{file_hash}_{clean_name}{ext}"
            
            # Limitar longitud
            if len(storage_name) > 200:
                # Truncar parte del nombre limpio
                max_clean_length = 200 - len(storage_name) + len(clean_name)
                clean_name = clean_name[:max_clean_length]
                storage_name = f"doc_{document_type_id}_{user_id}_{timestamp}_{file_hash}_{clean_name}{ext}"
            
            return storage_name
            
        except Exception as e:
            logger.warning(f"Error generando nombre de almacenamiento: {str(e)}")
            # Fallback
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            return f"doc_{timestamp}_{file_hash}{ext}"
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Limpiar nombre de archivo para almacenamiento seguro
        
        Args:
            filename: Nombre original
            
        Returns:
            str: Nombre limpio
        """
        try:
            # Caracteres prohibidos
            prohibited_chars = '<>:"/\\|?*'
            
            # Reemplazar caracteres prohibidos
            clean_name = filename
            for char in prohibited_chars:
                clean_name = clean_name.replace(char, '_')
            
            # Reemplazar espacios múltiples
            clean_name = ' '.join(clean_name.split())
            
            # Reemplazar espacios con guiones bajos
            clean_name = clean_name.replace(' ', '_')
            
            # Limitar longitud
            if len(clean_name) > 50:
                clean_name = clean_name[:50]
            
            # Asegurar que no esté vacío
            if not clean_name:
                clean_name = "documento"
            
            return clean_name.lower()
            
        except Exception:
            return "documento"
    
    def _calculate_file_hash(self, file_content: bytes) -> str:
        """
        Calcular hash SHA-256 del archivo
        
        Args:
            file_content: Contenido del archivo
            
        Returns:
            str: Hash SHA-256
        """
        try:
            sha256_hash = hashlib.sha256()
            sha256_hash.update(file_content)
            return sha256_hash.hexdigest()
        except Exception as e:
            logger.error(f"Error calculando hash: {str(e)}")
            return "unknown"
    
    async def _extract_file_metadata(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Extraer metadatos específicos del archivo
        
        Args:
            file_path: Ruta del archivo
            mime_type: Tipo MIME
            
        Returns:
            dict: Metadatos extraídos
        """
        try:
            metadata = {
                "extraction_timestamp": datetime.utcnow().isoformat(),
                "file_type": mime_type
            }
            
            # Metadatos para imágenes
            if mime_type.startswith('image/'):
                metadata.update(await self._extract_image_metadata(file_path))
            
            # Metadatos para PDFs
            elif mime_type == 'application/pdf':
                metadata.update(await self._extract_pdf_metadata(file_path))
            
            # Metadatos para documentos Word
            elif 'word' in mime_type or 'officedocument' in mime_type:
                metadata.update(await self._extract_document_metadata(file_path))
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos: {str(e)}")
            return {"error": str(e)}
    
    async def _extract_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extraer metadatos de imagen
        
        Args:
            file_path: Ruta de la imagen
            
        Returns:
            dict: Metadatos de imagen
        """
        try:
            with Image.open(file_path) as img:
                metadata = {
                    "width": img.width,
                    "height": img.height,
                    "format": img.format,
                    "mode": img.mode,
                    "has_transparency": img.mode in ('RGBA', 'LA', 'P')
                }
                
                # Extraer datos EXIF si existen
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = {}
                    exif = img._getexif()
                    
                    for tag_id, value in exif.items():
                        tag = ExifTags.TAGS.get(tag_id, tag_id)
                        
                        # Solo incluir metadatos útiles
                        if tag in ['DateTime', 'DateTimeOriginal', 'Software', 'Make', 'Model']:
                            try:
                                exif_data[tag] = str(value)
                            except:
                                pass
                    
                    if exif_data:
                        metadata["exif"] = exif_data
                
                return metadata
                
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos de imagen: {str(e)}")
            return {"error": str(e)}
    
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extraer metadatos de PDF
        
        Args:
            file_path: Ruta del PDF
            
        Returns:
            dict: Metadatos del PDF
        """
        try:
            doc = fitz.open(file_path)
            
            metadata = {
                "page_count": len(doc),
                "is_encrypted": doc.needs_pass,
                "has_toc": bool(doc.get_toc()),
                "pdf_version": doc.pdf_version() if hasattr(doc, 'pdf_version') else None
            }
            
            # Metadatos del documento
            doc_metadata = doc.metadata
            if doc_metadata:
                # Filtrar metadatos útiles
                useful_metadata = {}
                for key in ['title', 'author', 'subject', 'creator', 'producer', 'creationDate', 'modDate']:
                    if doc_metadata.get(key):
                        useful_metadata[key] = doc_metadata[key]
                
                if useful_metadata:
                    metadata["document_info"] = useful_metadata
            
            # Información de la primera página
            if len(doc) > 0:
                first_page = doc[0]
                rect = first_page.rect
                metadata["page_size"] = {
                    "width": rect.width,
                    "height": rect.height,
                    "unit": "points"
                }
            
            doc.close()
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos de PDF: {str(e)}")
            return {"error": str(e)}
    
    async def _extract_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extraer metadatos de documento Word
        
        Args:
            file_path: Ruta del documento
            
        Returns:
            dict: Metadatos del documento
        """
        try:
            # Para documentos Word, usar python-docx si está disponible
            try:
                from docx import Document
                
                doc = Document(file_path)
                
                metadata = {
                    "paragraph_count": len(doc.paragraphs),
                    "has_tables": len(doc.tables) > 0,
                    "table_count": len(doc.tables),
                    "has_images": bool([r for r in doc.part.rels.values() 
                                      if "image" in r.target_ref])
                }
                
                # Propiedades del documento
                core_props = doc.core_properties
                if core_props:
                    props = {}
                    for prop in ['title', 'author', 'subject', 'creator', 'created', 'modified']:
                        value = getattr(core_props, prop, None)
                        if value:
                            props[prop] = str(value)
                    
                    if props:
                        metadata["properties"] = props
                
                return metadata
                
            except ImportError:
                logger.warning("python-docx no disponible para extraer metadatos de Word")
                return {"info": "Metadatos de Word no disponibles"}
                
        except Exception as e:
            logger.warning(f"Error extrayendo metadatos de documento: {str(e)}")
            return {"error": str(e)}
    
    # === GESTIÓN DE ARCHIVOS ===
    
    def move_file(self, source_path: str, destination_path: str) -> bool:
        """
        Mover archivo de forma segura
        
        Args:
            source_path: Ruta origen
            destination_path: Ruta destino
            
        Returns:
            bool: True si se movió exitosamente
        """
        try:
            # Crear directorio destino si no existe
            os.makedirs(os.path.dirname(destination_path), exist_ok=True)
            
            # Mover archivo
            shutil.move(source_path, destination_path)
            
            logger.info(f"Archivo movido: {source_path} -> {destination_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error moviendo archivo: {str(e)}")
            return False
    
    def copy_file(self, source_path: str, destination_path: str) -> bool:
        """
        Copiar archivo de forma segura
        
        Args:
            source_path: Ruta origen
            destination_path: Ruta destino
            
        Returns:
            bool: True si se copió exitosamente
        """
        try:
            # Crear directorio destino si no existe
            os.makedirs(os.path.dirname(destination_path), exist_ok=True)
            
            # Copiar archivo
            shutil.copy2(source_path, destination_path)
            
            logger.info(f"Archivo copiado: {source_path} -> {destination_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error copiando archivo: {str(e)}")
            return False
    
    def delete_file(self, file_path: str) -> bool:
        """
        Eliminar archivo de forma segura
        
        Args:
            file_path: Ruta del archivo
            
        Returns:
            bool: True si se eliminó exitosamente
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Archivo eliminado: {file_path}")
                return True
            else:
                logger.warning(f"Archivo no encontrado para eliminar: {file_path}")
                return False
                
        except Exception as e:
            logger.error(f"Error eliminando archivo: {str(e)}")
            return False
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Obtener información de archivo
        
        Args:
            file_path: Ruta del archivo
            
        Returns:
            dict: Información del archivo
        """
        try:
            if not os.path.exists(file_path):
                return {"error": "Archivo no encontrado"}
            
            stat = os.stat(file_path)
            
            return {
                "size": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed": datetime.fromtimestamp(stat.st_atime).isoformat(),
                "is_file": os.path.isfile(file_path),
                "is_readable": os.access(file_path, os.R_OK),
                "is_writable": os.access(file_path, os.W_OK),
                "extension": os.path.splitext(file_path)[1].lower(),
                "filename": os.path.basename(file_path)
            }
            
        except Exception as e:
            logger.error(f"Error obteniendo información de archivo: {str(e)}")
            return {"error": str(e)}
    
    # === PROCESAMIENTO EN LOTE ===
    
    async def process_multiple_uploads(
        self, 
        files: List[UploadFile], 
        document_type_id: int,
        user_id: int
    ) -> List[Dict[str, Any]]:
        """
        Procesar múltiples archivos subidos
        
        Args:
            files: Lista de archivos
            document_type_id: ID del tipo de documento
            user_id: ID del usuario
            
        Returns:
            List[dict]: Resultados del procesamiento
        """
        try:
            logger.info(f"Procesando {len(files)} archivos en lote")
            
            results = []
            
            for i, file in enumerate(files):
                try:
                    result = await self.process_upload(file, document_type_id, user_id)
                    results.append({
                        "index": i,
                        "filename": file.filename,
                        "status": "success",
                        "result": result
                    })
                    
                except Exception as e:
                    logger.warning(f"Error procesando archivo {file.filename}: {str(e)}")
                    results.append({
                        "index": i,
                        "filename": file.filename,
                        "status": "error",
                        "error": str(e)
                    })
            
            successful = len([r for r in results if r["status"] == "success"])
            logger.info(f"Procesamiento en lote completado: {successful}/{len(files)} exitosos")
            
            return results
            
        except Exception as e:
            logger.error(f"Error en procesamiento en lote: {str(e)}")
            raise FileHandlerError(f"Error procesando archivos en lote: {str(e)}")
    
    # === LIMPIEZA Y MANTENIMIENTO ===
    
    def cleanup_temp_files(self, max_age_hours: int = 24):
        """
        Limpiar archivos temporales antiguos
        
        Args:
            max_age_hours: Edad máxima en horas
        """
        try:
            logger.info(f"Limpiando archivos temporales mayores a {max_age_hours} horas")
            
            current_time = datetime.now()
            deleted_count = 0
            total_size = 0
            
            for file_path in Path(self.temp_path).rglob("*"):
                if file_path.is_file():
                    try:
                        file_age = current_time - datetime.fromtimestamp(file_path.stat().st_mtime)
                        
                        if file_age.total_seconds() > (max_age_hours * 3600):
                            file_size = file_path.stat().st_size
                            file_path.unlink()
                            deleted_count += 1
                            total_size += file_size
                            
                    except Exception as e:
                        logger.warning(f"Error eliminando archivo temporal {file_path}: {str(e)}")
            
            size_mb = round(total_size / (1024 * 1024), 2)
            logger.info(f"Eliminados {deleted_count} archivos temporales ({size_mb} MB)")
            
        except Exception as e:
            logger.error(f"Error limpiando archivos temporales: {str(e)}")
    
    def get_storage_stats(self) -> Dict[str, Any]:
        """
        Obtener estadísticas de almacenamiento
        
        Returns:
            dict: Estadísticas de uso de almacenamiento
        """
        try:
            logger.info("Calculando estadísticas de almacenamiento")
            
            stats = {
                "documents_path": self.documents_path,
                "temp_path": self.temp_path,
                "timestamp": datetime.utcnow().isoformat()
            }
            
            # Estadísticas de documentos
            doc_stats = self._calculate_directory_stats(self.documents_path)
            stats["documents"] = doc_stats
            
            # Estadísticas de archivos temporales
            temp_stats = self._calculate_directory_stats(self.temp_path)
            stats["temp"] = temp_stats
            
            # Total
            stats["total"] = {
                "files": doc_stats["files"] + temp_stats["files"],
                "size_bytes": doc_stats["size_bytes"] + temp_stats["size_bytes"],
                "size_mb": round((doc_stats["size_bytes"] + temp_stats["size_bytes"]) / (1024 * 1024), 2)
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"Error calculando estadísticas: {str(e)}")
            return {"error": str(e)}
    
    def _calculate_directory_stats(self, directory: str) -> Dict[str, Any]:
        """
        Calcular estadísticas de un directorio
        
        Args:
            directory: Ruta del directorio
            
        Returns:
            dict: Estadísticas del directorio
        """
        try:
            total_size = 0
            file_count = 0
            
            for file_path in Path(directory).rglob("*"):
                if file_path.is_file():
                    file_count += 1
                    total_size += file_path.stat().st_size
            
            return {
                "files": file_count,
                "size_bytes": total_size,
                "size_mb": round(total_size / (1024 * 1024), 2),
                "size_gb": round(total_size / (1024 * 1024 * 1024), 2)
            }
            
        except Exception as e:
            logger.warning(f"Error calculando estadísticas de directorio: {str(e)}")
            return {"files": 0, "size_bytes": 0, "size_mb": 0, "size_gb": 0}


# === INSTANCIA GLOBAL ===

# Instancia singleton del manejador de archivos
_file_handler = None

def get_file_handler() -> FileHandler:
    """
    Obtener instancia del manejador de archivos
    
    Returns:
        FileHandler: Instancia del manejador
    """
    global _file_handler
    if _file_handler is None:
        _file_handler = FileHandler()
    return _file_handler


# === FUNCIONES DE UTILIDAD ===

async def validate_and_save_file(
    file: UploadFile,
    document_type_id: int,
    user_id: int,
    allowed_types: List[str] = None
) -> Dict[str, Any]:
    """
    Función de conveniencia para validar y guardar archivo
    
    Args:
        file: Archivo a procesar
        document_type_id: ID del tipo de documento
        user_id: ID del usuario
        allowed_types: Tipos MIME permitidos
        
    Returns:
        dict: Resultado del procesamiento
        
    Raises:
        HTTPException: Si la validación falla
    """
    try:
        handler = get_file_handler()
        
        # Validar archivo con tipos específicos si se proporcionan
        if allowed_types:
            is_valid, error_msg = handler.validate_file(file, allowed_types)
        else:
            is_valid, error_msg = handler.validate_file(file)
        
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error_msg
            )
        
        # Procesar archivo
        result = await handler.process_upload(file, document_type_id, user_id)
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en validate_and_save_file: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error procesando archivo"
        )


def get_safe_filename(filename: str) -> str:
    """
    Obtener nombre de archivo seguro
    
    Args:
        filename: Nombre original
        
    Returns:
        str: Nombre seguro
    """
    try:
        handler = get_file_handler()
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1]
        safe_base = handler._sanitize_filename(base_name)
        return f"{safe_base}{extension}"
    except:
        return "archivo_seguro.txt"


def calculate_file_hash(file_path: str) -> Optional[str]:
    """
    Calcular hash de archivo existente
    
    Args:
        file_path: Ruta del archivo
        
    Returns:
        Optional[str]: Hash SHA-256 o None si hay error
    """
    try:
        with open(file_path, 'rb') as f:
            sha256_hash = hashlib.sha256()
            for chunk in iter(lambda: f.read(4096), b""):
                sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
    except Exception as e:
        logger.error(f"Error calculando hash de archivo: {str(e)}")
        return None


def is_file_type_allowed(filename: str, mime_type: str = None) -> bool:
    """
    Verificar si un tipo de archivo está permitido
    
    Args:
        filename: Nombre del archivo
        mime_type: Tipo MIME (opcional)
        
    Returns:
        bool: True si está permitido
    """
    try:
        handler = get_file_handler()
        
        # Verificar extensión
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext not in handler.allowed_extensions:
            return False
        
        # Verificar tipo MIME si se proporciona
        if mime_type and mime_type not in handler.allowed_types:
            return False
        
        return True
        
    except:
        return False


def get_file_size_limit() -> int:
    """
    Obtener límite de tamaño de archivo en bytes
    
    Returns:
        int: Límite en bytes
    """
    return settings.MAX_FILE_SIZE


def get_file_size_limit_mb() -> int:
    """
    Obtener límite de tamaño de archivo en MB
    
    Returns:
        int: Límite en MB
    """
    return settings.MAX_FILE_SIZE // (1024 * 1024)


def create_temp_file(content: bytes, suffix: str = ".tmp") -> str:
    """
    Crear archivo temporal con contenido
    
    Args:
        content: Contenido del archivo
        suffix: Sufijo del archivo
        
    Returns:
        str: Ruta del archivo temporal
    """
    try:
        with tempfile.NamedTemporaryFile(
            mode='wb', 
            suffix=suffix, 
            dir=settings.TEMP_PATH,
            delete=False
        ) as temp_file:
            temp_file.write(content)
            return temp_file.name
    except Exception as e:
        logger.error(f"Error creando archivo temporal: {str(e)}")
        raise FileHandlerError(f"Error creando archivo temporal: {str(e)}")


def ensure_directory_exists(directory_path: str) -> bool:
    """
    Asegurar que un directorio existe
    
    Args:
        directory_path: Ruta del directorio
        
    Returns:
        bool: True si el directorio existe o se creó exitosamente
    """
    try:
        os.makedirs(directory_path, exist_ok=True)
        return True
    except Exception as e:
        logger.error(f"Error creando directorio {directory_path}: {str(e)}")
        return False


async def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Extraer texto de archivo (PDF, Word, etc.)
    
    Args:
        file_path: Ruta del archivo
        
    Returns:
        Optional[str]: Texto extraído o None
    """
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return await _extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await _extract_text_from_word(file_path)
        else:
            logger.warning(f"Extracción de texto no soportada para: {file_ext}")
            return None
            
    except Exception as e:
        logger.error(f"Error extrayendo texto de archivo: {str(e)}")
        return None


async def _extract_text_from_pdf(file_path: str) -> Optional[str]:
    """Extraer texto de PDF"""
    try:
        doc = fitz.open(file_path)
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_content.append(page.get_text())
        
        doc.close()
        return "\n".join(text_content)
        
    except Exception as e:
        logger.error(f"Error extrayendo texto de PDF: {str(e)}")
        return None


async def _extract_text_from_word(file_path: str) -> Optional[str]:
    """Extraer texto de documento Word"""
    try:
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            return "\n".join(text_content)
            
        except ImportError:
            logger.warning("python-docx no disponible para extraer texto de Word")
            return None
            
    except Exception as e:
        logger.error(f"Error extrayendo texto de Word: {str(e)}")
        return None


def check_disk_space(path: str = None) -> Dict[str, Any]:
    """
    Verificar espacio disponible en disco
    
    Args:
        path: Ruta a verificar (usa documents_path por defecto)
        
    Returns:
        dict: Información de espacio en disco
    """
    try:
        if not path:
            path = settings.DOCUMENTS_PATH
        
        statvfs = os.statvfs(path)
        
        # Calcular espacio
        total_space = statvfs.f_frsize * statvfs.f_blocks
        free_space = statvfs.f_frsize * statvfs.f_available
        used_space = total_space - free_space
        
        return {
            "path": path,
            "total_bytes": total_space,
            "used_bytes": used_space,
            "free_bytes": free_space,
            "total_gb": round(total_space / (1024**3), 2),
            "used_gb": round(used_space / (1024**3), 2),
            "free_gb": round(free_space / (1024**3), 2),
            "used_percent": round((used_space / total_space) * 100, 1),
            "free_percent": round((free_space / total_space) * 100, 1)
        }
        
    except Exception as e:
        logger.error(f"Error verificando espacio en disco: {str(e)}")
        return {"error": str(e)}


def cleanup_old_files(directory: str, max_age_days: int = 30) -> Dict[str, Any]:
    """
    Limpiar archivos antiguos de un directorio
    
    Args:
        directory: Directorio a limpiar
        max_age_days: Edad máxima en días
        
    Returns:
        dict: Resultado de la limpieza
    """
    try:
        logger.info(f"Limpiando archivos antiguos en {directory} (>{max_age_days} días)")
        
        current_time = datetime.now()
        cutoff_time = current_time - timedelta(days=max_age_days)
        
        deleted_count = 0
        total_size = 0
        errors = []
        
        for file_path in Path(directory).rglob("*"):
            if file_path.is_file():
                try:
                    file_mtime = datetime.fromtimestamp(file_path.stat().st_mtime)
                    
                    if file_mtime < cutoff_time:
                        file_size = file_path.stat().st_size
                        file_path.unlink()
                        deleted_count += 1
                        total_size += file_size
                        
                except Exception as e:
                    errors.append(f"Error eliminando {file_path}: {str(e)}")
        
        return {
            "directory": directory,
            "max_age_days": max_age_days,
            "deleted_files": deleted_count,
            "freed_bytes": total_size,
            "freed_mb": round(total_size / (1024**2), 2),
            "errors": errors,
            "timestamp": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error limpiando archivos antiguos: {str(e)}")
        return {"error": str(e)}


# === VALIDADORES PARA FASTAPI ===

def create_file_validator(max_size_mb: int = None, allowed_types: List[str] = None):
    """
    Crear validador de archivos personalizado
    
    Args:
        max_size_mb: Tamaño máximo en MB
        allowed_types: Tipos MIME permitidos
        
    Returns:
        function: Función validadora
    """
    def file_validator(file: UploadFile):
        """Validar archivo subido"""
        handler = get_file_handler()
        
        # Usar límites personalizados o por defecto
        if max_size_mb:
            max_bytes = max_size_mb * 1024 * 1024
            if hasattr(file, 'size') and file.size and file.size > max_bytes:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"Archivo demasiado grande. Máximo: {max_size_mb}MB"
                )
        
        # Validar archivo
        validation_types = allowed_types or handler.allowed_types
        is_valid, error_msg = handler.validate_file(file, validation_types)
        
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error_msg
            )
        
        return file
    
    return file_validator


# === HEALTH CHECK ===

async def file_handler_health_check() -> Dict[str, Any]:
    """
    Verificar estado del manejador de archivos
    
    Returns:
        dict: Estado del sistema de archivos
    """
    try:
        handler = get_file_handler()
        
        # Verificar directorios
        directories_status = {
            "documents_exists": os.path.exists(handler.documents_path),
            "documents_writable": os.access(handler.documents_path, os.W_OK),
            "temp_exists": os.path.exists(handler.temp_path),
            "temp_writable": os.access(handler.temp_path, os.W_OK)
        }
        
        # Verificar espacio en disco
        disk_info = check_disk_space()
        low_space = disk_info.get("free_percent", 100) < 10  # Menos del 10% libre
        
        # Obtener estadísticas
        storage_stats = handler.get_storage_stats()
        
        overall_status = (
            all(directories_status.values()) and 
            not low_space and
            "error" not in storage_stats
        )
        
        return {
            "status": "healthy" if overall_status else "unhealthy",
            "directories": directories_status,
            "disk_space": disk_info,
            "storage_stats": storage_stats,
            "warnings": ["Poco espacio en disco"] if low_space else [],
            "timestamp": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error en health check de file handler: {str(e)}")
        return {
            "status": "error",
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat()
        }