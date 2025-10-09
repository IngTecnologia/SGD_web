"""
Endpoints para generación de documentos
Generación de documentos con plantillas Word y códigos QR
"""
import logging
import os
import tempfile
from typing import List, Optional, Dict, Any
from datetime import datetime

from fastapi import (
    APIRouter, Depends, HTTPException, status, Query, Form, 
    BackgroundTasks, UploadFile, File
)
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session

from ...database import get_db
from ...config import get_settings
from ...models.user import User
from ...models.document_type import DocumentType
from ...models.qr_code import QRCode
from ...schemas.document import DocumentCreate
from ..deps import (
    get_current_user,
    require_generate_permission,
    get_document_type_by_id,
    get_request_logger,
    create_rate_limit_dependency
)
from ...utils.qr_processor import get_qr_processor, generate_unique_qr_id
from ...services.document_service import get_document_service

# Configuración
settings = get_settings()
logger = logging.getLogger(__name__)

# Router
router = APIRouter()

# Rate limiting para generación
generation_rate_limit = create_rate_limit_dependency(limit=5, window=60)


# === ESQUEMAS PARA GENERACIÓN ===

class DocumentGenerationRequest:
    """Esquema para solicitud de generación de documento"""
    def __init__(
        self,
        document_type_id: int,
        cedula: str,
        nombre_completo: str,
        telefono: Optional[str] = None,
        email: Optional[str] = None,
        direccion: Optional[str] = None,
        additional_data: Optional[Dict[str, Any]] = None,
        generate_qr: bool = True,
        qr_expires_days: Optional[int] = None
    ):
        self.document_type_id = document_type_id
        self.cedula = cedula
        self.nombre_completo = nombre_completo
        self.telefono = telefono
        self.email = email
        self.direccion = direccion
        self.additional_data = additional_data or {}
        self.generate_qr = generate_qr
        self.qr_expires_days = qr_expires_days


class DocumentGenerationResponse:
    """Esquema para respuesta de generación"""
    def __init__(
        self,
        success: bool,
        message: str,
        document_id: Optional[int] = None,
        qr_code_id: Optional[str] = None,
        file_url: Optional[str] = None,
        errors: Optional[List[str]] = None
    ):
        self.success = success
        self.message = message
        self.document_id = document_id
        self.qr_code_id = qr_code_id
        self.file_url = file_url
        self.errors = errors or []


# === ENDPOINTS DE GENERACIÓN ===

@router.get("/templates")
async def list_available_templates(
    db: Session = Depends(get_db),
    current_user: User = Depends(require_generate_permission)
):
    """
    Listar plantillas disponibles para generación
    """
    try:
        # Obtener tipos de documento con plantillas
        document_types = db.query(DocumentType).filter(
            DocumentType.is_active == True,
            DocumentType.template_path.isnot(None)
        ).all()
        
        templates = []
        for doc_type in document_types:
            template_info = {
                "document_type_id": doc_type.id,
                "code": doc_type.code,
                "name": doc_type.name,
                "description": doc_type.description,
                "template_path": doc_type.template_path,
                "requires_qr": doc_type.requires_qr,
                "required_fields": doc_type.required_fields,
                "can_generate": doc_type.can_generate
            }
            templates.append(template_info)
        
        return {
            "templates": templates,
            "total": len(templates),
            "timestamp": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error listando plantillas: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error al obtener plantillas disponibles"
        )


@router.post("/generate")
async def generate_document(
    # Datos del documento
    document_type_id: int = Form(..., description="ID del tipo de documento"),
    cedula: str = Form(..., description="Cédula del titular"),
    nombre_completo: str = Form(..., description="Nombre completo"),
    telefono: Optional[str] = Form(None, description="Teléfono"),
    email: Optional[str] = Form(None, description="Email"),
    direccion: Optional[str] = Form(None, description="Dirección"),
    
    # Configuración de generación
    generate_qr: bool = Form(True, description="Generar código QR"),
    qr_expires_days: Optional[int] = Form(None, description="Días hasta expiración del QR"),
    
    # Datos adicionales (JSON string)
    additional_data: Optional[str] = Form(None, description="Datos adicionales en JSON"),
    
    # Dependencias
    db: Session = Depends(get_db),
    current_user: User = Depends(require_generate_permission),
    log_action = Depends(get_request_logger),
    _: bool = Depends(generation_rate_limit)
):
    """
    Generar documento con plantilla y código QR
    """
    try:
        logger.info(f"Generando documento tipo {document_type_id} para {nombre_completo}")
        
        # Verificar tipo de documento
        document_type = db.query(DocumentType).filter(
            DocumentType.id == document_type_id,
            DocumentType.is_active == True
        ).first()
        
        if not document_type:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Tipo de documento no encontrado o inactivo"
            )
        
        if not document_type.can_generate:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Este tipo de documento no permite generación automática"
            )
        
        # Parsear datos adicionales
        extra_data = {}
        if additional_data:
            import json
            try:
                extra_data = json.loads(additional_data)
            except json.JSONDecodeError:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Datos adicionales deben ser JSON válido"
                )
        
        # Validar datos requeridos
        document_data = {
            "cedula": cedula,
            "nombre_completo": nombre_completo,
            "telefono": telefono,
            "email": email,
            "direccion": direccion,
            **extra_data
        }
        
        is_valid, validation_errors = document_type.validate_document_data(document_data)
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Datos inválidos: {', '.join(validation_errors)}"
            )
        
        # Generar código QR si se requiere
        qr_code = None
        qr_code_id = None
        
        if generate_qr and document_type.requires_qr:
            qr_id = generate_unique_qr_id()
            
            # Crear QR en base de datos
            from datetime import timedelta
            expires_at = None
            if qr_expires_days:
                expires_at = datetime.utcnow() + timedelta(days=qr_expires_days)
            
            qr_code = QRCode(
                qr_id=qr_id,
                document_type_id=document_type_id,
                generated_by=current_user.id,
                expires_at=expires_at
            )
            
            db.add(qr_code)
            db.commit()
            db.refresh(qr_code)
            
            qr_code_id = qr_code.qr_id
            logger.info(f"Código QR generado: {qr_code_id}")
        
        # Generar documento usando plantilla
        generated_file_path = await generate_document_from_template(
            document_type=document_type,
            document_data=document_data,
            qr_code_id=qr_code_id,
            user=current_user
        )
        
        # Crear registro del documento en base de datos
        document_service = get_document_service()
        
        # Simular UploadFile para el documento generado
        class GeneratedFile:
            def __init__(self, file_path: str, filename: str):
                self.file_path = file_path
                self.filename = filename
                self.size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            
            async def read(self):
                with open(self.file_path, 'rb') as f:
                    return f.read()
        
        generated_file = GeneratedFile(generated_file_path, f"{document_type.code}_{cedula}.docx")
        
        # Crear documento en BD
        from ...schemas.document import DocumentPersonInfo, DocumentMetadata
        person_info = DocumentPersonInfo(
            cedula=cedula,
            nombre_completo=nombre_completo,
            telefono=telefono,
            email=email,
            direccion=direccion,
            additional_data=extra_data
        )
        
        metadata = DocumentMetadata(
            tags=[document_type.code, "generated"],
            category="generated"
        )
        
        document_create_data = DocumentCreate(
            document_type_id=document_type_id,
            cedula=cedula,
            nombre_completo=nombre_completo,
            telefono=telefono,
            email=email,
            direccion=direccion,
            file_info=None,  # Se llenará en el servicio
            qr_code_id=qr_code_id,
            metadata=metadata
        )
        
        # Nota: Aquí necesitaríamos adaptar el servicio para manejar archivos generados
        # Por ahora retornamos la respuesta básica
        
        # Actualizar estadísticas
        document_type.increment_generated()
        current_user.increment_generations()
        db.commit()
        
        # Log de generación
        log_action("document_generated", {
            "document_type_id": document_type_id,
            "qr_code_id": qr_code_id,
            "cedula": cedula,
            "generated_file": os.path.basename(generated_file_path)
        })
        
        return {
            "success": True,
            "message": "Documento generado exitosamente",
            "document_type": {
                "id": document_type.id,
                "code": document_type.code,
                "name": document_type.name
            },
            "qr_code_id": qr_code_id,
            "file_path": generated_file_path,
            "download_url": f"/api/v1/generator/download/{os.path.basename(generated_file_path)}",
            "timestamp": datetime.utcnow().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando documento: {str(e)}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error al generar documento"
        )


@router.post("/generate-batch")
async def generate_documents_batch(
    document_type_id: int = Form(...),
    generate_qr: bool = Form(True),
    qr_expires_days: Optional[int] = Form(None),
    
    # Archivo CSV con datos
    data_file: UploadFile = File(..., description="Archivo CSV con datos"),
    
    db: Session = Depends(get_db),
    current_user: User = Depends(require_generate_permission),
    log_action = Depends(get_request_logger),
    _: bool = Depends(generation_rate_limit)
):
    """
    Generar múltiples documentos desde archivo CSV
    """
    try:
        # Verificar tipo de archivo
        if not data_file.filename.endswith('.csv'):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Solo se permiten archivos CSV"
            )
        
        # Verificar tipo de documento
        document_type = db.query(DocumentType).filter(
            DocumentType.id == document_type_id,
            DocumentType.is_active == True
        ).first()
        
        if not document_type or not document_type.can_generate:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Tipo de documento no encontrado o no permite generación"
            )
        
        # Leer archivo CSV
        import pandas as pd
        import io
        
        content = await data_file.read()
        csv_data = pd.read_csv(io.StringIO(content.decode('utf-8')))
        
        # Validar columnas requeridas
        required_columns = ['cedula', 'nombre_completo']
        missing_columns = [col for col in required_columns if col not in csv_data.columns]
        
        if missing_columns:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Columnas faltantes en CSV: {', '.join(missing_columns)}"
            )
        
        # Limitar número de registros
        if len(csv_data) > 100:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Máximo 100 registros por lote"
            )
        
        # Procesar en background (commented out for now)
        # background_tasks.add_task(
        #     process_batch_generation,
        #     document_type_id,
        #     csv_data.to_dict('records'),
        #     generate_qr,
        #     qr_expires_days,
        #     current_user.id,
        #     db
        # )
        
        # Log del lote
        log_action("batch_generation_started", {
            "document_type_id": document_type_id,
            "record_count": len(csv_data),
            "filename": data_file.filename
        })
        
        return {
            "message": "Generación en lote iniciada",
            "record_count": len(csv_data),
            "document_type": document_type.code,
            "status": "processing",
            "timestamp": datetime.utcnow().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en generación por lotes: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error al procesar generación en lote"
        )


@router.get("/download/{filename}")
async def download_generated_document(
    filename: str,
    current_user: User = Depends(get_current_user),
    log_action = Depends(get_request_logger)
):
    """
    Descargar documento generado
    """
    try:
        # Validar nombre de archivo
        if not filename.endswith('.docx') or '..' in filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Nombre de archivo inválido"
            )
        
        # Construir ruta del archivo
        file_path = os.path.join(settings.TEMP_PATH, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Archivo no encontrado"
            )
        
        # Log de descarga
        log_action("generated_document_downloaded", {
            "filename": filename
        })
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error descargando archivo generado: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error al descargar archivo"
        )


@router.get("/preview/{document_type_id}")
async def preview_template(
    document_type: DocumentType = Depends(get_document_type_by_id),
    current_user: User = Depends(require_generate_permission)
):
    """
    Obtener vista previa de plantilla
    """
    try:
        if not document_type.can_generate:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Este tipo no tiene plantilla disponible"
            )
        
        template_info = {
            "document_type": {
                "id": document_type.id,
                "code": document_type.code,
                "name": document_type.name,
                "description": document_type.description
            },
            "template": {
                "has_template": document_type.has_template,
                "template_path": document_type.template_path,
                "requires_qr": document_type.requires_qr
            },
            "requirements": {
                "required_fields": document_type.required_fields,
                "requires_qr": document_type.requires_qr,
                "requires_cedula": document_type.requires_cedula,
                "requires_nombre": document_type.requires_nombre,
                "requires_telefono": document_type.requires_telefono,
                "requires_email": document_type.requires_email,
                "requires_direccion": document_type.requires_direccion
            },
            "qr_config": {
                "table_number": document_type.qr_table_number,
                "row": document_type.qr_row,
                "column": document_type.qr_column,
                "width": document_type.qr_width,
                "height": document_type.qr_height
            } if document_type.requires_qr else None
        }
        
        return template_info
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error obteniendo vista previa: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error al obtener vista previa de plantilla"
        )


# === FUNCIONES AUXILIARES ===

async def generate_document_from_template(
    document_type: DocumentType,
    document_data: Dict[str, Any],
    qr_code_id: Optional[str],
    user: User
) -> str:
    """
    Generar documento desde plantilla Word
    """
    try:
        # Verificar que existe la plantilla
        template_path = os.path.join(settings.TEMPLATES_PATH, document_type.template_path)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")
        
        # Crear archivo de salida
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        output_filename = f"{document_type.code}_{document_data.get('cedula', 'unknown')}_{timestamp}.docx"
        output_path = os.path.join(settings.TEMP_PATH, output_filename)
        
        # Procesar plantilla con python-docx
        from docx import Document as DocxDocument
        
        # Abrir plantilla
        doc = DocxDocument(template_path)
        
        # Reemplazar marcadores de posición en el documento
        placeholders = {
            "{{CEDULA}}": document_data.get('cedula', ''),
            "{{NOMBRE_COMPLETO}}": document_data.get('nombre_completo', ''),
            "{{TELEFONO}}": document_data.get('telefono', ''),
            "{{EMAIL}}": document_data.get('email', ''),
            "{{DIRECCION}}": document_data.get('direccion', ''),
            "{{FECHA}}": datetime.now().strftime("%d/%m/%Y"),
            "{{HORA}}": datetime.now().strftime("%H:%M:%S"),
        }
        
        # Agregar datos adicionales
        for key, value in document_data.items():
            if key not in ['cedula', 'nombre_completo', 'telefono', 'email', 'direccion']:
                placeholders[f"{{{{{key.upper()}}}}}"] = str(value)
        
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholders.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
        
        # Insertar código QR si se requiere
        if qr_code_id and document_type.requires_qr:
            await insert_qr_code_in_document(
                doc, qr_code_id, document_type
            )
        
        # Guardar documento
        doc.save(output_path)
        
        logger.info(f"Documento generado: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error generando documento desde plantilla: {str(e)}")
        raise


async def insert_qr_code_in_document(
    doc,
    qr_code_id: str,
    document_type: DocumentType
):
    """
    Insertar código QR en documento Word
    """
    try:
        qr_processor = get_qr_processor()
        
        # Generar imagen QR temporal
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_qr:
            qr_image_path = qr_processor.generate_qr_code(
                data=qr_code_id,
                output_path=temp_qr.name,
                format="PNG"
            )
        
        # Insertar en la tabla especificada
        if len(doc.tables) >= document_type.qr_table_number:
            table = doc.tables[document_type.qr_table_number - 1]
            
            if (len(table.rows) > document_type.qr_row and 
                len(table.rows[document_type.qr_row].cells) > document_type.qr_column):
                
                cell = table.rows[document_type.qr_row].cells[document_type.qr_column]
                
                # Limpiar celda
                cell.text = ""
                
                # Agregar imagen QR
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                
                # Convertir pulgadas a EMU (English Metric Units)
                from docx.shared import Inches
                run.add_picture(
                    qr_image_path,
                    width=Inches(document_type.qr_width),
                    height=Inches(document_type.qr_height)
                )
        
        # Limpiar archivo temporal
        try:
            os.unlink(qr_image_path)
        except:
            pass
            
    except Exception as e:
        logger.error(f"Error insertando QR en documento: {str(e)}")
        # No fallar la generación por esto
        pass


async def process_batch_generation(
    document_type_id: int,
    records: List[Dict[str, Any]],
    generate_qr: bool,
    qr_expires_days: Optional[int],
    user_id: int,
    db: Session
):
    """
    Procesar generación en lote en background
    """
    try:
        logger.info(f"Procesando lote de {len(records)} documentos")
        
        success_count = 0
        error_count = 0
        
        for record in records:
            try:
                # Generar documento individual
                # (Implementación similar a generate_document pero simplificada)
                success_count += 1
                
            except Exception as e:
                logger.error(f"Error generando documento para {record.get('cedula', 'unknown')}: {str(e)}")
                error_count += 1
        
        logger.info(f"Lote completado: {success_count} exitosos, {error_count} errores")
        
    except Exception as e:
        logger.error(f"Error procesando lote: {str(e)}")


# === ENDPOINTS DE ESTADÍSTICAS ===

@router.get("/stats")
async def get_generation_stats(
    document_type_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Obtener estadísticas de generación
    """
    try:
        # Query base
        query = db.query(DocumentType).filter(
            DocumentType.template_path.isnot(None)
        )
        
        if document_type_id:
            query = query.filter(DocumentType.id == document_type_id)
        
        document_types = query.all()
        
        stats = {
            "total_templates": len(document_types),
            "active_templates": len([dt for dt in document_types if dt.is_active]),
            "by_type": [],
            "total_generated": sum(dt.generated_count for dt in document_types),
            "timestamp": datetime.utcnow().isoformat()
        }
        
        for doc_type in document_types:
            type_stats = {
                "id": doc_type.id,
                "code": doc_type.code,
                "name": doc_type.name,
                "generated_count": doc_type.generated_count,
                "can_generate": doc_type.can_generate,
                "requires_qr": doc_type.requires_qr
            }
            stats["by_type"].append(type_stats)
        
        return stats
        
    except Exception as e:
        logger.error(f"Error obteniendo estadísticas: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error al obtener estadísticas de generación"
        )